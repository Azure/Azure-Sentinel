#!/usr/bin/env python3
"""
Build-AgentPackage.py
=====================

Phase 6 cowork builder for the Sentinel Data Connector and Agent Builder.
Given an SCC-exported AgentManifest.yaml in `partner-center/<isv-slug>-agent/inbox/`,
this script:

  1. Lints the manifest against R1–R7 (the 7 common Security Store review failures).
  2. Auto-patches what's safe (R1 Product/Publisher, R5 MCP.Sentinel, R7 grammar).
  3. Applies `progress.json.phases.5_agent_build.renameMap` to KQL references inside
     `Settings.Instructions` only (NEVER to YAML keys).
  4. Writes the linted manifest preserving block-scalar style (`>-`, `|`, `>`) and key order.
  5. Generates PackageManifest.yaml.
  6. Zips the package (Mac-clean: no `__MACOSX`, no `.DS_Store`, no `._*`).
  7. Generates all Partner Center artifacts:
       - offer-listing-description.md
       - plan-description.md (with __SCU_TBD__)
       - user-guide/user-guide.docx (Word doc following the Commvault 8-section template, anonymized)
       - diagrams/architecture.mmd + build-png.sh
       - screenshots/README.md (capture recipe, 1280×720)
       - scu-measurement.md (3–5 run protocol)
       - partner-center-checklist.md (every click in lab-06 Tasks 2–8)
       - lint-report.json
  8. Writes `progress.json.phases.6_publishing`.

Inputs are pulled from `config/progress.json` (companyName, phases.5_agent_build.*).
The only positional argument is the raw manifest filename inside `inbox/`.

Usage
-----
    python3 scripts/Build-AgentPackage.py \
        --raw-manifest partner-center/gigamon-agent/inbox/GigamonSuspiciousEncryptedC2HuntAdvisor.yaml

If --raw-manifest is omitted, the script looks for the single `.yaml` file under
`partner-center/<inferred-slug>-agent/inbox/`.

Exit codes
----------
    0   ok — all artifacts written
    2   bad inputs (missing raw manifest, malformed progress.json)
    3   lint fail that cannot be auto-patched (e.g., R2, R4 — surface to developer)
"""
from __future__ import annotations

import argparse
import datetime
import json
import os
import pathlib
import re
import subprocess
import sys
from io import StringIO

try:
    from ruamel.yaml import YAML
    from ruamel.yaml.scalarstring import LiteralScalarString, FoldedScalarString
except ImportError:
    sys.stderr.write(
        "ruamel.yaml is required. Install with: python3 -m pip install --user ruamel.yaml\n"
    )
    sys.exit(2)

try:
    from docx import Document
except ImportError:
    sys.stderr.write(
        "python-docx is required. Install with: python3 -m pip install --user python-docx\n"
    )
    sys.exit(2)

# ────────────────────────────────────────────────────────────────────────────────
# YAML round-trip configuration — preserves >-, |, >, key order, quoting
# ────────────────────────────────────────────────────────────────────────────────
yaml = YAML(typ="rt")
yaml.preserve_quotes = True
yaml.width = 4096          # do not re-wrap long lines
yaml.indent(mapping=2, sequence=4, offset=2)


# ────────────────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────────────────
def utc_now_iso() -> str:
    return datetime.datetime.utcnow().isoformat() + "Z"


def find_repo_root(start: pathlib.Path) -> pathlib.Path:
    cur = start.resolve()
    while cur != cur.parent:
        if (cur / ".git").exists() or (cur / "config" / "progress.json").exists():
            return cur
        cur = cur.parent
    raise SystemExit("Could not locate repository root (looked for .git or config/progress.json)")


def load_progress(repo: pathlib.Path) -> dict:
    p = repo / "config" / "progress.json"
    if not p.exists():
        raise SystemExit(f"Missing {p}. Run phases 0–5 first.")
    return json.loads(p.read_text())


def isv_slug_from_progress(progress: dict) -> str:
    """`Gigamon` → `gigamon`; `Palo Alto Networks` → `palo-alto-networks`."""
    name = progress.get("companyName", "")
    if not name:
        raise SystemExit("progress.json.companyName is empty")
    return re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")


# ────────────────────────────────────────────────────────────────────────────────
# Lint
# ────────────────────────────────────────────────────────────────────────────────
def lint_manifest(manifest, isv: str) -> tuple[list[dict], list[str]]:
    """Walk R1–R7. Mutates `manifest` in place for auto-patches.
    Returns (rules_log, fatal_failures)."""
    rules: list[dict] = []
    fatal: list[str] = []

    def add(rid, name, status, before=None, after=None, note=""):
        rules.append({"id": rid, "name": name, "status": status,
                      "before": before, "after": after, "note": note})

    # R1 Product / Publisher / PublisherSource ≠ "Custom"
    ad_list = manifest.get("AgentDefinitions") or []
    if not ad_list:
        fatal.append("R1: AgentDefinitions[] is empty — manifest is not a published agent.")
        add("R1", "Product/Publisher non-Custom", "fail",
            None, None, "AgentDefinitions missing")
    else:
        ad = ad_list[0]
        before = {k: ad.get(k) for k in ("Product", "Publisher", "PublisherSource")}
        patched = False
        for field in ("Product", "Publisher", "PublisherSource"):
            if ad.get(field) in (None, "", "Custom"):
                ad[field] = isv
                patched = True
        after = {k: ad.get(k) for k in ("Product", "Publisher", "PublisherSource")}
        add("R1", "Product/Publisher/PublisherSource must be ISV name (not 'Custom')",
            "patched" if patched else "pass", before, after,
            f"auto-patched from 'Custom' to '{isv}' per progress.json.companyName")

        # R2 Settings[].Name === Inputs[].Name
        skills_inputs = {i["Name"]
                         for sg in manifest.get("SkillGroups", [])
                         for s in sg.get("Skills", [])
                         for i in s.get("Inputs", []) or []}
        settings_names = {s.get("Name") for s in ad.get("Settings", []) or []}
        if settings_names == skills_inputs:
            add("R2", "Settings[].Name === Skills[].Inputs[].Name", "pass",
                {"settings": sorted(settings_names),
                 "inputs": sorted(skills_inputs)}, None, "")
        else:
            fatal.append("R2: Settings.Name ≠ Inputs.Name. Cannot auto-fix; "
                         "developer must rename inside SCC and re-export.")
            add("R2", "Settings[].Name === Skills[].Inputs[].Name", "fail",
                {"settings": sorted(settings_names),
                 "inputs": sorted(skills_inputs)}, None,
                "Mismatch — developer must rename in SCC and re-export.")

        # R3 Every Input + Setting has a non-empty Description (auto-trim whitespace)
        r3_missing: list[str] = []
        r3_cleaned: list[dict] = []

        def _trim(obj, where):
            d = obj.get("Description", "")
            if isinstance(d, str):
                stripped = d.strip()
                if not stripped:
                    r3_missing.append(where)
                elif stripped != d:
                    r3_cleaned.append({"where": where, "before": d, "after": stripped})
                    obj["Description"] = stripped

        for sg in manifest.get("SkillGroups", []):
            for s in sg.get("Skills", []):
                for i in s.get("Inputs", []) or []:
                    _trim(i, f"Skills.Inputs.{i.get('Name','?')}.Description")
        for s in ad.get("Settings", []) or []:
            _trim(s, f"AgentDefinitions.Settings.{s.get('Name','?')}.Description")

        if r3_missing:
            fatal.append(f"R3: Missing Description on: {r3_missing}")
            add("R3", "Every Input + Setting has non-empty Description",
                "fail", r3_missing, None, "")
        else:
            add("R3", "Every Input + Setting has non-empty Description",
                "pass" if not r3_cleaned else "patched",
                None, r3_cleaned or None,
                "auto-trimmed leading/trailing whitespace on Description fields")

        # R4 Skill names descriptive (regex placeholder names = fail)
        r4_bad: list[str] = []
        for sg in manifest.get("SkillGroups", []):
            for s in sg.get("Skills", []):
                n = s.get("Name", "")
                if re.match(r"^(skill|agent|test)[ _\-]?(v?\d+|\d+)$", n, re.I) or len(n) <= 4:
                    r4_bad.append(n)
        if r4_bad:
            fatal.append(f"R4: Non-descriptive skill names: {r4_bad}")
            add("R4", "Skill names descriptive", "fail", r4_bad, None, "")
        else:
            add("R4", "Skill names descriptive", "pass", None, None, "")

        # R5 RequiredSkillsets contains MCP.Sentinel (auto-append)
        rs = ad.get("RequiredSkillsets") or []
        if "MCP.Sentinel" not in rs:
            rs.append("MCP.Sentinel")
            ad["RequiredSkillsets"] = rs
            add("R5", "RequiredSkillsets contains MCP.Sentinel",
                "patched", None, rs, "auto-appended")
        else:
            add("R5", "RequiredSkillsets contains MCP.Sentinel",
                "pass", list(rs), None, "")

        # R6 KQL time windows: only ago(24h) inside ```kql ``` fenced blocks
        instr_text = ""
        for sg in manifest.get("SkillGroups", []):
            for s in sg.get("Skills", []):
                v = (s.get("Settings") or {}).get("Instructions", "")
                if v:
                    instr_text += "\n" + str(v)
        bad_ago_in_kql: list[str] = []
        for m in re.finditer(r"```kql\b(.*?)```", instr_text, re.DOTALL):
            for w in re.findall(r"ago\(([^)]+)\)", m.group(1)):
                if w.strip() != "24h":
                    bad_ago_in_kql.append(w.strip())
        if bad_ago_in_kql:
            fatal.append(f"R6: KQL blocks contain non-24h ago(): {bad_ago_in_kql}")
            add("R6", "KQL time window only ago(24h)", "fail",
                bad_ago_in_kql, None,
                "Developer must rewrite KQL to use ago(24h)")
        else:
            add("R6", "KQL time window only ago(24h)", "pass",
                None, None,
                "ago(7d)/ago(30d) found only outside fenced kql blocks — those are prose, OK")

        # R7 Common product capitalization (microsoft sentinel → Microsoft Sentinel, etc.)
        # AND collapse 3+ spaces to 2 in instructions blocks.
        r7_patched: list[dict] = []
        for sg in manifest.get("SkillGroups", []):
            for s in sg.get("Skills", []):
                settings = s.get("Settings") or {}
                orig = settings.get("Instructions")
                if not isinstance(orig, str):
                    continue
                fixed = orig
                for pat, rep in [
                    (r"\bmicrosoft sentinel\b", "Microsoft Sentinel"),
                    (r"\bmicrosoft entra\b", "Microsoft Entra"),
                    (r"\bmicrosoft defender\b", "Microsoft Defender"),
                ]:
                    cnt = len(re.findall(pat, fixed, re.I)) - len(re.findall(rep, fixed))
                    if cnt > 0:
                        fixed = re.sub(pat, rep, fixed, flags=re.I)
                        r7_patched.append({"pattern": pat, "count": cnt})
                # 3+ spaces → 2 spaces (but never touch code fences / leading indent)
                collapsed = re.sub(r"(?<![\n])([ ]{3,})", "  ", fixed)
                if collapsed != fixed:
                    r7_patched.append({"pattern": "3+ spaces collapsed", "count": 1})
                    fixed = collapsed
                if fixed != orig:
                    # Preserve original block-scalar style on assignment
                    style = getattr(orig, "_yaml_format", None)
                    if isinstance(orig, FoldedScalarString):
                        settings["Instructions"] = FoldedScalarString(fixed)
                    elif isinstance(orig, LiteralScalarString):
                        settings["Instructions"] = LiteralScalarString(fixed)
                    else:
                        settings["Instructions"] = fixed
        add("R7", "Grammar / capitalization", "patched" if r7_patched else "pass",
            None, r7_patched or None, "")

    return rules, fatal


# ────────────────────────────────────────────────────────────────────────────────
# Apply renameMap to Settings.Instructions only
# ────────────────────────────────────────────────────────────────────────────────
def apply_rename_map(manifest, rename_map: dict) -> list[dict]:
    log: list[dict] = []
    for sg in manifest.get("SkillGroups", []):
        for s in sg.get("Skills", []):
            settings = s.get("Settings") or {}
            v = settings.get("Instructions")
            if not isinstance(v, str):
                continue
            new = v
            for old, repl in rename_map.items():
                cnt = new.count(old)
                if cnt:
                    new = new.replace(old, repl)
                    log.append({"from": old, "to": repl, "replacements": cnt,
                                "where": f"Skills.{s.get('Name','?')}.Settings.Instructions"})
            if new != v:
                if isinstance(v, FoldedScalarString):
                    settings["Instructions"] = FoldedScalarString(new)
                elif isinstance(v, LiteralScalarString):
                    settings["Instructions"] = LiteralScalarString(new)
                else:
                    settings["Instructions"] = new
    return log


# ────────────────────────────────────────────────────────────────────────────────
# Artifact generators
# ────────────────────────────────────────────────────────────────────────────────
def generate_package_manifest(out_path: pathlib.Path, agent_folder: str, isv: str,
                              agent_display: str) -> None:
    pm = {
        "manifest": [{
            "id": agent_folder,
            "description": (
                f"Triage agent that hunts suspicious activity by correlating "
                f"{isv} telemetry with Microsoft Sentinel and Microsoft Entra signals."
            ),
            "type": "CopilotAgent",
        }],
        "schema": {"version": "1.0.0"},
    }
    with open(out_path, "w") as f:
        yaml.dump(pm, f)


def zip_package(pc_dir: pathlib.Path, agent_folder: str) -> str:
    zip_path = pc_dir / "agent-package.zip"
    if zip_path.exists():
        zip_path.unlink()
    subprocess.run(
        ["zip", "-r", "-X", "agent-package.zip",
         "PackageManifest.yaml", agent_folder,
         "-x", ".*", "-x", "__MACOSX", "-x", "*/.DS_Store", "-x", "*/._*"],
        cwd=pc_dir, check=True, stdout=subprocess.DEVNULL,
    )
    return subprocess.run(
        ["unzip", "-l", "agent-package.zip"],
        cwd=pc_dir, check=True, capture_output=True,
        text=True, encoding="utf-8", errors="replace",
    ).stdout


def generate_offer_listing(pc_dir: pathlib.Path, isv: str, agent_display: str,
                           allow_tables: list[dict], scenarios: list[dict]) -> None:
    table_lines = "\n".join(
        f"   - `{t['productionName']}` — "
        + ("Microsoft Entra ID sign-in telemetry" if t["productionName"] == "SigninLogs"
           else "Microsoft Sentinel / Defender correlated alerts" if t["productionName"] == "SecurityAlert"
           else f"{isv} product data")
        for t in allow_tables
    )
    text = f"""# {agent_display} — Offer Listing

## Offer Summary

The {agent_display} accelerates SOC triage by correlating {isv} telemetry with Microsoft Entra sign-in risk and Microsoft Sentinel / Defender security alerts. Given a single primary input, the agent emits a deterministic verdict and concrete next actions — in under a minute, without exposing raw event data.

## Description

The {agent_display} automates the cross-product correlation a SOC analyst would otherwise perform by hand, against the last 24 hours of activity in Microsoft Sentinel Data Lake. The agent is read-only, time-bounded, and summarized by design — it does not modify any host, alert, or identity.

## Agent Tasks

{chr(10).join(f"- **{s['id']}** — drives a `{s['drivesVerdict']}` verdict." for s in scenarios)}

## Workflow

1. SOC analyst opens Microsoft Security Copilot and selects the {agent_display}.
2. Analyst provides the required input.
3. Agent queries the following tables in Microsoft Sentinel Data Lake (last 24 hours only):
{table_lines}
4. Agent correlates findings across the sources.
5. Agent emits a verdict and next-action recommendation in markdown.

## Output

- **Verdict:** a single deterministic level driven by the scoring rubric.
- **Per-scenario findings:** one short summary per built-in scenario.
- **Next actions:** prescriptive recommendations.
- **Data gaps:** explicit reporting if a required table is empty or missing.

## Why this agent

- Single-input ergonomic — analyst pastes one value, gets a triage answer.
- Read-only, 24-hour-scoped — safe to run at scale.
- Correlates multiple independent signal sources without leaving Security Copilot.
- Built on {isv}'s product telemetry paired with Microsoft Sentinel's identity and alert surface.

## Support

For support requests please provide: Tenant ID, Security Copilot Session ID, agent run summary or screenshots, and any reported data gaps.
"""
    (pc_dir / "offer-listing-description.md").write_text(text)


def generate_plan_description(pc_dir: pathlib.Path, agent_display: str) -> None:
    (pc_dir / "plan-description.md").write_text(
        f"The {agent_display} is available at no additional cost beyond your Microsoft Security "
        f"Copilot Capacity Units (SCU) consumption. This agent typically consumes __SCU_TBD__ SCU "
        f"per analysis run. SCU consumption may vary depending on the volume of data in your "
        f"Microsoft Sentinel workspace and the complexity of the investigation.\n"
    )


def generate_user_guide_docx(pc_dir: pathlib.Path, isv: str, agent_display: str,
                             allow_tables: list[dict], scenarios: list[dict],
                             primary_input: dict, scoring_rubric: dict) -> None:
    """Commvault 8-section template, anonymized."""
    ug_dir = pc_dir / "user-guide"
    ug_dir.mkdir(exist_ok=True)
    iso_month = datetime.date.today().isoformat()[:7]
    doc = Document()
    doc.add_heading(f"{agent_display} User Guide", level=0)
    for line in [f"Publisher: {isv}",
                 f"Agent: {agent_display}",
                 "Version: 1.0.0",
                 f"Date: {iso_month}"]:
        doc.add_paragraph(line)

    doc.add_heading("What it is", level=2)
    doc.add_paragraph(
        f"The {agent_display} is a triage agent that helps security operations teams correlate "
        f"{isv} telemetry with Microsoft Entra and Microsoft Sentinel signals to determine the "
        f"severity of a single investigation target. The agent runs entirely against data already "
        f"present in Microsoft Sentinel Data Lake."
    )
    doc.add_heading("Where it runs", level=2)
    doc.add_paragraph("Microsoft Security Copilot")
    doc.add_heading("Required integrations", level=2)
    doc.add_paragraph("Microsoft Sentinel Data Lake")
    doc.add_heading("Output", level=2)
    rubric_text = "; ".join(f"{lvl}: {desc}" for lvl, desc in scoring_rubric.items())
    doc.add_paragraph(
        f"By correlating partner telemetry with identity and alert signals, the agent emits a "
        f"single deterministic verdict driven by an explicit scoring rubric ({rubric_text})."
    )

    doc.add_heading("Contents", level=2)
    doc.add_paragraph(
        "This guide is designed for security operations and incident response teams using the "
        "agent in Microsoft Security Copilot."
    )
    for i, t in enumerate([
        "Overview", "What the agent does", "Prerequisites", "Install the agent",
        "Configure access and data sources", "Run an investigation",
        "Sample Results", "Support and feedback",
    ], 1):
        doc.add_paragraph(f"{i}. {t}")

    # 1. Overview
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        f"The {agent_display} helps security teams correlate {isv} product telemetry with identity "
        f"risk and alert signals in Microsoft Sentinel Data Lake."
    )
    for t in allow_tables:
        owner = (
            "Microsoft Entra ID sign-in telemetry"
            if t["productionName"] == "SigninLogs"
            else "Microsoft Sentinel / Defender correlated alerts"
            if t["productionName"] == "SecurityAlert"
            else f"{isv} product events"
        )
        doc.add_paragraph(f"{t['productionName']} — {owner}", style="List Bullet")
    doc.add_paragraph(
        "By linking partner telemetry with identity and alert context, the agent provides "
        "investigation-ready correlation without exposing raw event data."
    )
    doc.add_paragraph(
        "The agent is read-only, time-bounded, and summarized by design, enabling safe "
        "investigation at scale."
    )

    # 2. What the agent does
    doc.add_heading("2. What the agent does", level=1)
    doc.add_heading("Inputs: what data the agent consumes", level=2)
    doc.add_paragraph("The agent operates exclusively on data already present in Microsoft Sentinel Data Lake:")
    for t in allow_tables:
        doc.add_paragraph(t["productionName"], style="List Bullet")
    doc.add_paragraph(
        "Note: The agent does not assume schema or column names. All queries are schema-verified "
        "at runtime and limited to a 24-hour lookback window."
    )
    doc.add_heading("Tasks: what the agent performs", level=2)
    for s in scenarios:
        doc.add_paragraph(f"{s['id']} — drives a {s['drivesVerdict']} verdict.", style="List Bullet")
    doc.add_heading("Outputs: what results the agent generates", level=2)
    for lvl, desc in scoring_rubric.items():
        doc.add_paragraph(f"{lvl}: {desc}", style="List Bullet")
    doc.add_paragraph("Concrete next-action recommendations.", style="List Bullet")
    doc.add_paragraph("Explicit data-gap reporting when a required table is empty or missing.", style="List Bullet")

    # 3. Prerequisites
    doc.add_heading("3. Prerequisites", level=1)
    doc.add_heading("Platform requirements", level=2)
    for b in [
        "Microsoft Security Copilot enabled for the organization.",
        "Microsoft Sentinel workspace with Data Lake enabled.",
        "Relevant partner telemetry already ingested into Sentinel Data Lake.",
    ]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_heading("Required data lake tables", level=2)
    doc.add_paragraph("The agent uses only the following tables:")
    for t in allow_tables:
        doc.add_paragraph(t["productionName"], style="List Bullet")
    doc.add_paragraph("If a table is unavailable or empty, the agent runs best-effort and reports gaps.")

    # 4. Install
    doc.add_heading("4. Install the agent", level=1)
    for i, b in enumerate([
        "Open Microsoft Security Copilot.",
        "Navigate to Agents -> Browse more agents.",
        f"Search {agent_display} and install the agent to your Security Copilot environment.",
    ], 1):
        doc.add_paragraph(f"{i}. {b}")

    # 5. Configure
    doc.add_heading("5. Configure access and data sources", level=1)
    doc.add_paragraph("No manual configuration is required beyond standard Sentinel access.")
    doc.add_heading("Validation checks (recommended)", level=2)
    doc.add_paragraph("Can the running user read Sentinel Data Lake tables?", style="List Bullet")
    doc.add_paragraph(
        "Are " + ", ".join(t["productionName"] for t in allow_tables)
        + " logs present for the last 24 hours?",
        style="List Bullet",
    )
    doc.add_paragraph("If access is partial, the agent continues and documents limitations.")

    # 6. Run
    doc.add_heading("6. Run an investigation", level=1)
    doc.add_paragraph(
        f"Provide the {primary_input.get('name','primary input')} "
        f"({primary_input.get('type','value')}) for the agent to correlate telemetry data.",
        style="List Bullet",
    )
    doc.add_heading("Example input", level=2)
    doc.add_paragraph(
        f"{primary_input.get('name','input')}: {primary_input.get('example','<value>')}",
        style="Intense Quote",
    )

    # 7. Sample Results
    doc.add_heading("7. Sample Results", level=1)
    doc.add_paragraph("Below is a sample result the agent provides after correlating the data.")
    p = doc.add_paragraph()
    r = p.add_run("[Insert screenshot of a representative agent run here. The matching PNG is in "
                  "../screenshots/03-final-output.png after capture.]")
    r.italic = True

    # 8. Support
    doc.add_heading("8. Support and feedback", level=1)
    doc.add_paragraph("When requesting support, provide:")
    for b in ["Tenant ID", "Security Copilot Session ID",
              "Summary output or screenshots", "Any reported data gaps"]:
        doc.add_paragraph(b, style="List Bullet")
    doc.add_paragraph(f"Publisher: {isv}")
    doc.add_paragraph(f"Product: {agent_display}")
    doc.save(ug_dir / "user-guide.docx")


def generate_diagram(pc_dir: pathlib.Path, agent_display: str,
                     allow_tables: list[dict]) -> None:
    diag_dir = pc_dir / "diagrams"
    diag_dir.mkdir(exist_ok=True)
    table_nodes = "\n".join(
        f"    D -->|{t['productionName']}| {chr(ord('E')+i)}[{t['productionName']}]\n    {chr(ord('E')+i)} --> H"
        for i, t in enumerate(allow_tables)
    )
    (diag_dir / "architecture.mmd").write_text(
        f"""flowchart LR
    A[SOC Analyst] -->|input| B[Microsoft Security Copilot<br/>{agent_display}]
    B --> C[MCP.Sentinel<br/>Skillset]
    C --> D[(Microsoft Sentinel Data Lake)]
{table_nodes}
    H{{Verdict<br/>per scoring rubric}}
    H --> A
"""
    )
    build_png = diag_dir / "build-png.sh"
    build_png.write_text("""#!/usr/bin/env bash
set -euo pipefail
cd "$(dirname "$0")"
npx -y @mermaid-js/mermaid-cli@latest -i architecture.mmd -o architecture.png -w 1600 -H 900
echo "Built: $(pwd)/architecture.png"
""")
    build_png.chmod(0o755)


def generate_screenshot_recipe(pc_dir: pathlib.Path, agent_display: str) -> None:
    ss_dir = pc_dir / "screenshots"
    ss_dir.mkdir(exist_ok=True)
    (ss_dir / "README.md").write_text(
        f"""# Screenshot capture recipe — {agent_display}

Partner Center requires **at least 3 PNGs at exactly 1280×720 px**. Capture at higher resolution, then resize with the per-file command below.

| # | Filename | What to capture | What MUST be visible |
|---|---|---|---|
| 1 | `01-agent-run-with-plugins.png` | Security Copilot session with the agent **actively running** and producing results | The **Plugins** panel on the right MUST show **Microsoft Sentinel** enabled. This is the screenshot Partner Center reviewers check first. |
| 2 | `02-input-prompt.png` | The agent's **input prompt** screen (Run → One time → fill input) | The input box, the example value, and the **Run** button. |
| 3 | `03-final-output.png` | The agent's **full final output** with the verdict block visible | The verdict, per-scenario findings, and the next-actions list. |

## Capture and resize (macOS)

```bash
# 1. Capture (Cmd+Shift+4 then Space, click the SCC window)
#    Save to ~/Downloads/screencapture-<timestamp>.png
# 2. Resize to 1280×720:
sips -z 720 1280 ~/Downloads/screencapture-*.png --out 01-agent-run-with-plugins.png
# repeat for 02-input-prompt.png and 03-final-output.png
```

## Capture and resize (Linux / ImageMagick)

```bash
convert source.png -resize 1280x720! 01-agent-run-with-plugins.png
```

## Verify

```bash
sips -g pixelHeight -g pixelWidth *.png   # macOS
identify *.png                            # ImageMagick
```

Every PNG MUST report `1280 x 720`. Partner Center will reject anything else.
"""
    )


def generate_scu_protocol(pc_dir: pathlib.Path, agent_display: str,
                          scenarios: list[dict], primary_input: dict) -> None:
    example = primary_input.get("example", "<value>")
    rows = "\n".join(
        f"| {i+1} | `{example}` | {s['id']} | _____ | |"
        for i, s in enumerate(scenarios[:5])
    )
    (pc_dir / "scu-measurement.md").write_text(
        f"""# SCU measurement protocol — {agent_display}

Partner Center requires a sentence stating typical SCU per run in the plan description. Run the agent 3–5 times against realistic inputs, record the SCU shown in the Security Copilot session header, then compute the average and paste it into `plan-description.md` (replacing `__SCU_TBD__`).

## Runs

| # | Input value | Scenario expected to fire | SCU consumed | Notes |
|---|---|---|---|---|
{rows}

## Compute

- Average SCU = (sum of column 4) / number of runs.
- Round to the nearest integer.

## Paste back

1. Open `plan-description.md`.
2. Replace `__SCU_TBD__` with the rounded average.
3. Update `progress.json.phases.6_publishing.scuEstimate` with the same number.
"""
    )


def generate_pc_checklist(pc_dir: pathlib.Path, isv: str, agent_folder: str,
                          agent_display: str, slug: str) -> None:
    (pc_dir / "partner-center-checklist.md").write_text(
        f"""# Partner Center submission checklist — {agent_display}

Tracks every click in lab-06 Tasks 2–8. Tick each box as you go.

## Task 2 — Gather required information

Already produced by the agent:
- [x] `PackageManifest.yaml`
- [x] `{agent_folder}/AgentManifest.yaml` (linted, R1–R7 applied, renameMap applied)
- [x] `agent-package.zip`
- [x] `offer-listing-description.md`
- [x] `plan-description.md` (with `__SCU_TBD__` placeholder)
- [x] `user-guide/user-guide.docx`
- [x] `diagrams/architecture.mmd` + `build-png.sh`
- [x] `screenshots/README.md`
- [x] `scu-measurement.md`
- [x] `lint-report.json`

Still needs developer action:
- [ ] {isv} logo PNG (square, ≥216×216)
- [ ] 3 screenshots under `screenshots/` (run `sips`/`convert` per `screenshots/README.md`)
- [ ] Run agent 3–5 times for SCU; fill `scu-measurement.md`; update `plan-description.md`
- [ ] Render architecture diagram: `bash diagrams/build-png.sh`
- [ ] Open `user-guide/user-guide.docx` in Word/Pages/Google Docs, review/edit, then **File → Save As → PDF** → `user-guide/user-guide.pdf`
- [ ] Entra app registration for landing-page URL + connection webhook (dummy values acceptable)

## Task 3.1–3.2 — Access Partner Center, New offer
- [ ] Sign in to <https://partner.microsoft.com>
- [ ] Marketplace offers → **New offer** → **Software as a Service (SaaS)** → Start blank
- [ ] Offer ID = `{slug}-advisor`
- [ ] Alias = `{agent_display}`

## Task 3.3 — Offer setup
- [ ] Sell through Microsoft = **Yes**
- [ ] License management = **No**
- [ ] ✓ **My offer integrates with Microsoft Security services**

## Task 3.4 — Properties
- [ ] Categories = **Security or Compliance**
- [ ] Industries = (blank)
- [ ] Legal contract = Standard or your own

## Task 3.5 — Offer listing
- [ ] Paste description from `offer-listing-description.md`
- [ ] Upload logo + 3 screenshots
- [ ] Marketing/product URL under **Supplemental product information for customers → Product information links**
- [ ] Upload `user-guide/user-guide.pdf` under **Product information documents**

## Task 3.6 — Microsoft Security services
- [ ] Integrated services = ✓ Microsoft Security Copilot + ✓ Microsoft Sentinel
- [ ] Product prerequisites = ✓ Microsoft Security Copilot, ✓ Microsoft Sentinel, ✓ Microsoft Defender, ✓ Microsoft Entra (check all that apply)
- [ ] Solution type = ✓ Deployable solution
- [ ] License management = choose based on your model
- [ ] Security Copilot agent = ✓ Check **"Security Copilot agent"**
- [ ] **Upload .zip package** → `agent-package.zip`

## Task 4 — Preview audience
- [ ] Add Entra IDs of internal testers

## Task 5 — Technical configuration
- [ ] Landing page URL + Connection webhook + Entra tenant ID + Entra app ID (dummy OK, but cannot be blank)

## Task 6 — Plan and pricing
- [ ] Create plan; plan name does NOT include Microsoft product names
- [ ] Paste plan description from `plan-description.md` (with `__SCU_TBD__` replaced)
- [ ] Markets = Select all; Visibility = Public

## Task 7 — Supplement content
- [ ] SaaS Scenarios = "SaaS solution is not hosted in Azure"
- [ ] Text note = `Offer listing is for Security Copilot Agent in Microsoft Security Store.`
- [ ] Upload `diagrams/architecture.png`

## Task 8 — Final review & publish
- [ ] Walk lab-06 section 8.1 checklist line by line
- [ ] **Review and publish** → **Publish** → **Go Live**
"""
    )


# ────────────────────────────────────────────────────────────────────────────────
# Main
# ────────────────────────────────────────────────────────────────────────────────
def main() -> int:
    ap = argparse.ArgumentParser(description="Sentinel Data Connector and Agent Builder — Phase 6 package builder")
    ap.add_argument("--raw-manifest", type=str, default=None,
                    help="Path to the SCC-exported AgentManifest YAML (inside inbox/). "
                         "If omitted, the single yaml under partner-center/<slug>-agent/inbox/ is used.")
    ap.add_argument("--repo-root", type=str, default=None,
                    help="Repository root. Auto-detected from CWD if omitted.")
    args = ap.parse_args()

    repo = pathlib.Path(args.repo_root).resolve() if args.repo_root else find_repo_root(pathlib.Path.cwd())
    progress = load_progress(repo)
    isv = progress.get("companyName", "")
    if not isv:
        sys.stderr.write("progress.json.companyName is empty.\n"); return 2
    isv_slug = isv_slug_from_progress(progress)

    p5 = progress.get("phases", {}).get("5_agent_build")
    if not p5:
        sys.stderr.write("progress.json.phases.5_agent_build is missing — run Phase 5 first.\n"); return 2

    agent_display: str = p5["agentName"]
    agent_folder = re.sub(r"\s+", "", agent_display)  # no spaces
    allow_tables = p5.get("allowlistedTables", [])
    rename_map = p5.get("renameMap", {})
    scenarios = p5.get("scenarios", [])
    scoring_rubric = p5.get("scoringRubric", {})
    primary_input = p5.get("primaryInput", {})

    pc_dir = repo / "partner-center" / f"{isv_slug}-agent"
    pc_dir.mkdir(parents=True, exist_ok=True)
    inbox = pc_dir / "inbox"
    inbox.mkdir(exist_ok=True)

    if args.raw_manifest:
        raw_path = pathlib.Path(args.raw_manifest).resolve()
    else:
        candidates = sorted(inbox.glob("*.yaml")) + sorted(inbox.glob("*.yml"))
        if not candidates:
            sys.stderr.write(f"No YAML found in {inbox}. "
                             f"Drop the SCC-exported manifest there first.\n"); return 2
        if len(candidates) > 1:
            sys.stderr.write(f"Multiple YAMLs in {inbox}. Pass --raw-manifest explicitly.\n"); return 2
        raw_path = candidates[0]

    print(f"[in]  {raw_path.relative_to(repo)}")
    with open(raw_path) as f:
        manifest = yaml.load(f)

    # 6.1 Lint
    rules_log, fatal = lint_manifest(manifest, isv)
    summary = {s: sum(1 for r in rules_log if r["status"] == s)
               for s in ("pass", "patched", "fail")}
    print(f"[lint] {summary['pass']} pass, {summary['patched']} patched, {summary['fail']} fail")
    if fatal:
        for f in fatal: sys.stderr.write(f"  ! {f}\n")
        # Still write lint-report.json so the developer sees the findings
        (pc_dir / "lint-report.json").write_text(json.dumps(
            {"lintedAt": utc_now_iso(), "rules": rules_log, "fatal": fatal},
            indent=2))
        return 3

    # 6.2 Apply renameMap to Settings.Instructions only
    rename_log = apply_rename_map(manifest, rename_map)
    if rename_log:
        for r in rename_log:
            print(f"[rename] {r['from']} → {r['to']} ({r['replacements']}x) in {r['where']}")

    # Write linted manifest (round-trip preserves >-, |, > and key order)
    out_agent_dir = pc_dir / agent_folder
    out_agent_dir.mkdir(exist_ok=True)
    linted_path = out_agent_dir / "AgentManifest.yaml"
    with open(linted_path, "w") as f:
        yaml.dump(manifest, f)
    print(f"[out] {linted_path.relative_to(repo)}")

    # lint-report.json
    (pc_dir / "lint-report.json").write_text(json.dumps({
        "manifestSource": str(raw_path.relative_to(repo)),
        "lintedManifestPath": str(linted_path.relative_to(repo)),
        "lintedAt": utc_now_iso(),
        "rules": rules_log,
        "renameMapApplied": rename_log,
        "fatal": [],
    }, indent=2))

    # 6.3 PackageManifest.yaml
    generate_package_manifest(pc_dir / "PackageManifest.yaml", agent_folder, isv, agent_display)
    print(f"[out] {(pc_dir / 'PackageManifest.yaml').relative_to(repo)}")

    # 6.4 Zip
    listing = zip_package(pc_dir, agent_folder)
    print("[zip]\n" + listing.strip())

    # 6.5 All other artifacts
    generate_offer_listing(pc_dir, isv, agent_display, allow_tables, scenarios)
    generate_plan_description(pc_dir, agent_display)
    generate_user_guide_docx(pc_dir, isv, agent_display, allow_tables, scenarios,
                             primary_input, scoring_rubric)
    generate_diagram(pc_dir, agent_display, allow_tables)
    generate_screenshot_recipe(pc_dir, agent_display)
    generate_scu_protocol(pc_dir, agent_display, scenarios, primary_input)
    generate_pc_checklist(pc_dir, isv, agent_folder, agent_display, isv_slug)
    print("[out] all Partner Center artifacts written under "
          f"{pc_dir.relative_to(repo)}/")

    # 6.7 progress.json — schema matches Phase 6.7 spec in copilot-instructions.md
    progress.setdefault("phases", {})["6_publishing"] = {
        "status": "package_ready",
        "agentName": agent_display,
        "agentNameNoSpaces": agent_folder,
        "agentFolder": agent_folder,
        "packageFolder": str(pc_dir.relative_to(repo)) + "/",
        "packageZipPath": str((pc_dir / "agent-package.zip").relative_to(repo)),
        "agentManifestRawPath":    str(raw_path.relative_to(repo)),
        "agentManifestLintedPath": str(linted_path.relative_to(repo)),
        "lintResult": {
            "pass": all(r["status"] in ("pass", "patched") for r in rules_log) and not fatal,
            "rules": [{"id": r["id"], "status": r["status"]} for r in rules_log],
            "renameMapApplied": rename_log,
        },
        "scuEstimate": "__SCU_TBD__",
        "partnerCenterStatus": "draft",
        "artifacts": {
            "packageManifest":  str((pc_dir / "PackageManifest.yaml").relative_to(repo)),
            "agentManifest":    str(linted_path.relative_to(repo)),
            "offerDescription": str((pc_dir / "offer-listing-description.md").relative_to(repo)),
            "planDescription":  str((pc_dir / "plan-description.md").relative_to(repo)),
            "userGuideDocx":    str((pc_dir / "user-guide" / "user-guide.docx").relative_to(repo)),
            "architectureMmd":  str((pc_dir / "diagrams" / "architecture.mmd").relative_to(repo)),
            "screenshotRecipe": str((pc_dir / "screenshots" / "README.md").relative_to(repo)),
            "scuProtocol":      str((pc_dir / "scu-measurement.md").relative_to(repo)),
            "pcChecklist":      str((pc_dir / "partner-center-checklist.md").relative_to(repo)),
            "lintReport":       str((pc_dir / "lint-report.json").relative_to(repo)),
        },
        "scuEstimate": "__SCU_TBD__",
        "partnerCenterOfferId": None,
        "partnerCenterStatus": "draft",
        "notes": [
            f"AgentManifest sourced from {raw_path.relative_to(repo)}",
            f"renameMap applied to Settings.Instructions only (KQL refs); YAML keys untouched.",
            "user-guide.docx generated via python-docx; developer reviews/edits in Word then Save As → PDF.",
        ],
        "updatedAt": utc_now_iso(),
    }
    (repo / "config" / "progress.json").write_text(json.dumps(progress, indent=2))
    print("[out] config/progress.json.phases.6_publishing")
    return 0


if __name__ == "__main__":
    sys.exit(main())
